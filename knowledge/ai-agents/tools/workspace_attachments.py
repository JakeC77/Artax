"""Tools for querying and reading workspace scratchpad attachments."""

from __future__ import annotations

import asyncio
import io
import json
import logging
import os
from typing import Any, Optional
from urllib import error, request

from pydantic_ai import RunContext

from app.core.authenticated_graphql_client import execute_graphql, run_graphql
from app.tools import register_tool

logger = logging.getLogger(__name__)

# Get base URL from GraphQL endpoint (same base)
_DEFAULT_ENDPOINT = "https://app-sbx-westus3-01.azurewebsites.net/gql/"
_GRAPHQL_ENDPOINT = os.getenv("WORKSPACE_GRAPHQL_ENDPOINT", _DEFAULT_ENDPOINT)
# Extract base URL (remove /gql/ suffix)
_BASE_URL = _GRAPHQL_ENDPOINT.rstrip("/gql/").rstrip("/")
_REQUEST_TIMEOUT = float(os.getenv("WORKSPACE_GRAPHQL_TIMEOUT", "10"))

# Module-level cache for downloaded files (attachment_id -> bytes)
_download_cache: dict[str, bytes] = {}
_metadata_cache: dict[str, dict[str, Any]] = {}

_SCRATCHPAD_ATTACHMENTS_QUERY = """
query ScratchpadAttachments($workspaceId: UUID!) {
  scratchpadAttachments(where: { workspaceId: { eq: $workspaceId }}) {
    scratchpadAttachmentId
    workspaceId
    title
    description
    uri
    size
    fileType
    createdOn
  }
}
""".strip()


async def _run_graphql(query: str, variables: dict[str, Any], tenant_id: Optional[str] = None) -> dict[str, Any]:
    """Run the blocking GraphQL call in a background thread with authentication support."""
    return await run_graphql(query, variables, graphql_endpoint=_GRAPHQL_ENDPOINT, tenant_id=tenant_id)


def _download_attachment(attachment_id: str, tenant_id: str) -> bytes:
    """Download attachment file from API endpoint.
    
    Args:
        attachment_id: UUID of the attachment
        tenant_id: Tenant ID for the download URL
        
    Returns:
        File contents as bytes
        
    Raises:
        RuntimeError: If download fails
    """
    # Construct download URL
    download_url = f"{_BASE_URL}/scratchpad/attachments/{attachment_id}/download?tid={tenant_id}"
    
    logger.debug(f"Downloading attachment {attachment_id} from {download_url}")
    
    try:
        req = request.Request(download_url, method="GET")
        with request.urlopen(req, timeout=_REQUEST_TIMEOUT * 3) as resp:  # Longer timeout for file downloads
            file_bytes = resp.read()
        logger.debug(f"Downloaded {len(file_bytes)} bytes for attachment {attachment_id}")
        return file_bytes
    except error.HTTPError as exc:
        body = exc.read().decode("utf-8") if exc.fp else ""
        raise RuntimeError(
            f"Failed to download attachment {attachment_id}: HTTP {exc.code} - {body or exc.reason}"
        ) from exc
    except error.URLError as exc:
        raise RuntimeError(f"Failed to download attachment {attachment_id}: {exc.reason}") from exc


def _parse_page_range(page_range: str | None, total_pages: int) -> list[int]:
    """Parse page range string into list of page numbers.
    
    Args:
        page_range: String like "1-5", "1,3,5", "all", or None
        total_pages: Total number of pages in document
        
    Returns:
        List of 1-indexed page numbers to extract
    """
    if not page_range or page_range.lower() == "all":
        return list(range(1, total_pages + 1))
    
    pages = []
    parts = page_range.replace(" ", "").split(",")
    
    for part in parts:
        if "-" in part:
            # Range like "1-5"
            start_str, end_str = part.split("-", 1)
            try:
                start = int(start_str)
                end = int(end_str) if end_str else total_pages
                pages.extend(range(max(1, start), min(end + 1, total_pages + 1)))
            except ValueError:
                raise ValueError(f"Invalid page range format: {part}")
        else:
            # Single page number
            try:
                page_num = int(part)
                if 1 <= page_num <= total_pages:
                    pages.append(page_num)
            except ValueError:
                raise ValueError(f"Invalid page number: {part}")
    
    # Remove duplicates and sort
    return sorted(list(set(pages)))


async def _read_pdf(file_bytes: bytes, page_range: str | None = None) -> dict[str, Any]:
    """Extract text from PDF file.
    
    Args:
        file_bytes: PDF file contents
        page_range: Optional page range to extract (e.g., "1-5", "1,3,5", "all"). 
                    If None and document > 10 pages, defaults to first 5 pages.
    
    Returns:
        Dict with content, page_count, pages_extracted, and chunking info
    """
    try:
        import PyPDF2
    except ImportError:
        raise RuntimeError(
            "PyPDF2 is required for PDF reading. Install with: pip install PyPDF2"
        )
    
    try:
        pdf_file = io.BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        total_pages = len(pdf_reader.pages)
        
        # Auto-chunking: If document is large (>5 pages) and no range specified, 
        # default to first 3 pages (more conservative to avoid token limits)
        if page_range is None and total_pages > 5:
            page_range = "1-3"
            logger.info(f"Large PDF detected ({total_pages} pages). Extracting first 3 pages. "
                       f"Use page_range parameter to read specific pages.")
        
        # Parse page range
        pages_to_extract = _parse_page_range(page_range, total_pages)
        
        text_parts = []
        for page_num in pages_to_extract:
            try:
                page = pdf_reader.pages[page_num - 1]  # Convert to 0-indexed
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"--- Page {page_num} ---\n{text}")
            except Exception as e:
                logger.warning(f"Error extracting text from PDF page {page_num}: {e}")
                continue
        
        content = "\n\n".join(text_parts) if text_parts else "[PDF file contains no extractable text]"
        
        result = {
            "content": content,
            "page_count": total_pages,
            "pages_extracted": pages_to_extract,
            "pages_extracted_count": len(pages_to_extract),
        }
        
        # Add chunking guidance if not all pages were extracted
        if len(pages_to_extract) < total_pages:
            remaining_pages = total_pages - len(pages_to_extract)
            result["chunking_info"] = {
                "is_chunked": True,
                "total_pages": total_pages,
                "extracted_pages": f"{min(pages_to_extract)}-{max(pages_to_extract)}" if pages_to_extract else "none",
                "remaining_pages": remaining_pages,
                "suggestion": f"To read more pages, call again with page_range parameter (e.g., '6-10', '11-15', etc.)"
            }
        else:
            result["chunking_info"] = {
                "is_chunked": False,
                "total_pages": total_pages,
            }
        
        return result
    except Exception as e:
        raise RuntimeError(f"Error reading PDF: {e}") from e


async def _read_word(file_bytes: bytes, paragraph_range: str | None = None) -> dict[str, Any]:
    """Extract text from Word document.
    
    Args:
        file_bytes: Word document contents
        paragraph_range: Optional paragraph range (e.g., "1-50", "1,3,5", "all").
                        For large documents, can be used to extract specific sections.
    
    Returns:
        Dict with content, paragraph_count, and chunking info
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx is required for Word document reading. Install with: pip install python-docx"
        )
    
    try:
        doc_file = io.BytesIO(file_bytes)
        doc = Document(doc_file)
        
        # Extract all paragraphs
        all_paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        total_paragraphs = len(all_paragraphs)
        
        # Auto-chunking: If document is large (>50 paragraphs) and no range specified,
        # default to first 30 paragraphs (more conservative to avoid token limits)
        paragraphs_to_extract = None
        if paragraph_range is None and total_paragraphs > 50:
            paragraph_range = "1-30"
            logger.info(f"Large Word document detected ({total_paragraphs} paragraphs). "
                       f"Extracting first 30 paragraphs. Use paragraph_range parameter for specific sections.")
        
        if paragraph_range and paragraph_range.lower() != "all":
            # Parse paragraph range (similar to page range)
            paragraphs_to_extract = _parse_page_range(paragraph_range, total_paragraphs)
            paragraphs = [all_paragraphs[i - 1] for i in paragraphs_to_extract if 1 <= i <= total_paragraphs]
        else:
            paragraphs = all_paragraphs
            paragraphs_to_extract = list(range(1, total_paragraphs + 1))
        
        # Extract text from tables (always extract all tables for now)
        table_texts = []
        for table_idx, table in enumerate(doc.tables, 1):
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                if any(row_cells):  # Only include non-empty rows
                    table_rows.append(" | ".join(row_cells))
            if table_rows:
                table_texts.append(f"--- Table {table_idx} ---\n" + "\n".join(table_rows))
        
        content_parts = []
        if paragraphs:
            content_parts.append("\n".join(paragraphs))
        if table_texts:
            content_parts.extend(table_texts)
        
        content = "\n\n".join(content_parts) if content_parts else "[Word document contains no text]"
        
        result = {
            "content": content,
            "paragraph_count": total_paragraphs,
            "paragraphs_extracted": len(paragraphs),
            "table_count": len(doc.tables),
        }
        
        # Add chunking guidance if not all paragraphs were extracted
        if paragraphs_to_extract and len(paragraphs_to_extract) < total_paragraphs:
            remaining = total_paragraphs - len(paragraphs_to_extract)
            result["chunking_info"] = {
                "is_chunked": True,
                "total_paragraphs": total_paragraphs,
                "extracted_paragraphs": f"{min(paragraphs_to_extract)}-{max(paragraphs_to_extract)}" if paragraphs_to_extract else "none",
                "remaining_paragraphs": remaining,
                "suggestion": f"To read more paragraphs, call again with paragraph_range parameter (e.g., '51-100', etc.)"
            }
        else:
            result["chunking_info"] = {
                "is_chunked": False,
                "total_paragraphs": total_paragraphs,
            }
        
        return result
    except Exception as e:
        raise RuntimeError(f"Error reading Word document: {e}") from e


async def _read_csv(file_bytes: bytes, extract_tables: bool = False) -> dict[str, Any]:
    """Read CSV file and optionally extract structured data.
    
    Uses pandas if available (better formatting), falls back to built-in csv module.
    """
    csv_file = io.BytesIO(file_bytes)
    
    # Try pandas first (better formatting and table extraction)
    try:
        import pandas as pd
        
        try:
            df = pd.read_csv(csv_file)
            
            # Convert to text representation
            content = df.to_string(index=False)
            
            result = {
                "content": content,
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": list(df.columns),
            }
            
            # If extract_tables is True, also return structured data
            if extract_tables:
                # Convert DataFrame to list of dicts (one per row)
                result["table_data"] = df.to_dict("records")
            
            return result
        except Exception as e:
            raise RuntimeError(f"Error reading CSV with pandas: {e}") from e
            
    except ImportError:
        # Fallback to built-in csv module
        import csv
        
        logger.warning(
            "pandas not available, using built-in csv module. "
            "For better CSV handling, install pandas: pip install pandas"
        )
        
        try:
            # Reset file pointer
            csv_file.seek(0)
            
            # Read CSV with built-in module
            text_content = csv_file.read().decode("utf-8")
            reader = csv.DictReader(text_content.splitlines())
            
            rows = list(reader)
            columns = reader.fieldnames or []
            
            # Format as text table (simple formatting)
            if rows:
                # Header
                content_lines = [" | ".join(columns)]
                content_lines.append("-" * (sum(len(str(c)) for c in columns) + len(columns) * 3))
                
                # Rows
                for row in rows:
                    content_lines.append(" | ".join(str(row.get(col, "")) for col in columns))
                
                content = "\n".join(content_lines)
            else:
                content = " | ".join(columns) if columns else "[Empty CSV file]"
            
            result = {
                "content": content,
                "row_count": len(rows),
                "column_count": len(columns),
                "columns": list(columns),
            }
            
            # If extract_tables is True, return structured data
            if extract_tables:
                result["table_data"] = rows
            
            return result
        except Exception as e:
            raise RuntimeError(f"Error reading CSV: {e}") from e


async def _read_txt(file_bytes: bytes) -> dict[str, Any]:
    """Read plain text file."""
    try:
        # Try UTF-8 first, fallback to latin-1
        try:
            content = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            content = file_bytes.decode("latin-1")
        
        return {
            "content": content,
            "line_count": len(content.splitlines()),
        }
    except Exception as e:
        raise RuntimeError(f"Error reading text file: {e}") from e


def _normalize_file_type(file_type: Optional[str]) -> str:
    """Normalize file type from MIME type or extension to standard extension.
    
    Handles:
    - MIME types: "application/pdf" -> "pdf"
    - Extensions: ".pdf", "pdf" -> "pdf"
    - Case insensitive
    
    Args:
        file_type: File type string (MIME type or extension)
        
    Returns:
        Normalized file extension (e.g., "pdf", "docx", "csv")
    """
    if not file_type:
        return ""
    
    file_type = file_type.lower().strip()
    
    # Remove leading dot if present
    file_type = file_type.lstrip(".")
    
    # MIME type to extension mapping
    mime_to_ext = {
        "application/pdf": "pdf",
        "application/msword": "doc",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/csv": "csv",
        "text/plain": "txt",
        "text/txt": "txt",
    }
    
    # Check if it's a MIME type
    if file_type in mime_to_ext:
        return mime_to_ext[file_type]
    
    # Check if it contains a MIME type pattern
    for mime, ext in mime_to_ext.items():
        if mime in file_type or file_type in mime:
            return ext
    
    # Already an extension, return as-is
    return file_type


async def _read_file_content(
    file_bytes: bytes, 
    file_type: Optional[str], 
    extract_tables: bool = False,
    page_range: str | None = None,
) -> dict[str, Any]:
    """Read file content based on file type.
    
    Args:
        file_bytes: File contents as bytes
        file_type: File type/extension or MIME type (e.g., "pdf", "application/pdf", "docx")
        extract_tables: For CSV/Word, also extract structured table data
        page_range: For PDFs: page range to extract (e.g., "1-5", "all").
                   For Word docs: treated as paragraph_range.
        
    Returns:
        Dict with "content" (text) and metadata
    """
    # Normalize file_type (handle MIME types and extensions)
    normalized_type = _normalize_file_type(file_type)
    
    logger.debug(f"File type: '{file_type}' -> normalized: '{normalized_type}'")
    
    # Determine file type from extension or content
    if normalized_type in ["pdf"]:
        return await _read_pdf(file_bytes, page_range=page_range)
    elif normalized_type in ["docx", "doc"]:
        return await _read_word(file_bytes, paragraph_range=page_range)
    elif normalized_type in ["csv"]:
        return await _read_csv(file_bytes, extract_tables=extract_tables)
    elif normalized_type in ["txt", "text"]:
        return await _read_txt(file_bytes)
    else:
        # Try to detect from content or default to text
        if file_type:
            logger.warning(f"Unsupported file type: {file_type} (normalized: {normalized_type}). Attempting to read as text.")
        return await _read_txt(file_bytes)


@register_tool("scratchpad_attachments_list")
async def scratchpad_attachments_list(
    ctx: RunContext[dict],
    workspace_id: str | None = None,
    file_type: str | None = None,
) -> dict[str, Any]:
    """List scratchpad attachments for a workspace.
    
    Fetches attachment metadata from GraphQL including title, description, file type, size, etc.
    Use this to discover available attachments before reading them.
    
    Args:
        workspace_id: Workspace UUID. If not provided, will use workspace_id from context.
        file_type: Optional filter by file type (e.g., "pdf", "docx", "csv"). Case-insensitive.
    
    Returns:
        Dict with workspace_id and list of attachment metadata:
        {
            "workspace_id": "...",
            "attachments": [
                {
                    "scratchpadAttachmentId": "...",
                    "title": "...",
                    "description": "...",
                    "fileType": "...",
                    "size": 12345,
                    "createdOn": "...",
                    ...
                },
                ...
            ]
        }
    """
    tool_name = "scratchpad_attachments_list"
    logger.info(f"Tool called: {tool_name}")
    
    # Get workspace_id from context if not provided
    if not workspace_id:
        workspace_id = ctx.deps.get("workspace_id")
        if workspace_id:
            logger.debug(f"Using workspace_id from context: {workspace_id}")
        else:
            available_keys = list(ctx.deps.keys()) if ctx.deps else []
            logger.error(f"Tool '{tool_name}' failed: workspace_id not found in context. Available keys: {available_keys}")
            raise ValueError(
                f"workspace_id is required. Either pass it as a parameter or ensure it's in the workflow context. "
                f"Available context keys: {available_keys}"
            )
    
    tenant_id = ctx.deps.get("tenant_id")
    logger.info(f"Tool '{tool_name}' executing with workspace_id={workspace_id}, tenant_id={'present' if tenant_id else 'missing'}, file_type={file_type}")
    
    try:
        logger.debug(f"Fetching scratchpad attachments for workspace_id: {workspace_id}")
        data = await _run_graphql(_SCRATCHPAD_ATTACHMENTS_QUERY, {"workspaceId": workspace_id}, tenant_id=tenant_id)
        
        attachments = data.get("scratchpadAttachments", [])
        
        # Filter by file_type if provided
        if file_type:
            file_type_lower = file_type.lower().lstrip(".")
            attachments = [
                att for att in attachments
                if att.get("fileType", "").lower().lstrip(".") == file_type_lower
            ]
        
        # Cache metadata
        for att in attachments:
            att_id = att.get("scratchpadAttachmentId")
            if att_id:
                _metadata_cache[att_id] = att
        
        logger.info(f"Tool '{tool_name}' succeeded: found {len(attachments)} attachments")
        return {
            "workspace_id": workspace_id,
            "attachments": attachments,
            "count": len(attachments),
        }
    except Exception as e:
        logger.error(f"Tool '{tool_name}' failed with error: {type(e).__name__}: {str(e)}")
        raise


@register_tool("scratchpad_attachment_read")
async def scratchpad_attachment_read(
    ctx: RunContext[dict],
    attachment_id: str,
    extract_tables: bool = False,
    page_range: str | None = None,
) -> dict[str, Any]:
    """Download and read a scratchpad attachment file.
    
    Automatically handles:
    - Downloading the file (with caching to avoid re-downloads)
    - Detecting file type
    - Extracting text content
    - Chunking large documents (auto-chunks PDFs >10 pages, Word docs >100 paragraphs)
    - Optionally extracting structured data (tables from CSV/Word)
    
    Args:
        attachment_id: UUID of the attachment to read
        extract_tables: If True, also extract structured table data (for CSV/Word files)
        page_range: For PDFs: page range to extract (e.g., "1-5", "6-10", "1,3,5", "all").
                   For Word docs: paragraph range (e.g., "1-50", "51-100").
                   If None, auto-chunks large documents (first 5 pages for PDFs, 
                   first 50 paragraphs for Word docs).
    
    Returns:
        Dict with attachment metadata and extracted content:
        {
            "attachment_id": "...",
            "title": "...",
            "file_type": "pdf",
            "content": "extracted text...",
            "metadata": {
                "size": 12345,
                "page_count": 20,  # total pages
                "pages_extracted": [1, 2, 3, 4, 5],  # which pages were extracted
                "pages_extracted_count": 5,
                ...
            },
            "chunking_info": {
                "is_chunked": true,
                "total_pages": 20,
                "extracted_pages": "1-5",
                "remaining_pages": 15,
                "suggestion": "To read more pages, call again with page_range='6-10', etc."
            },
            "table_data": [...]  # if extract_tables=True and file has tables
        }
    
    Examples:
        # Read first 5 pages of a PDF
        scratchpad_attachment_read(attachment_id="...", page_range="1-5")
        
        # Read specific pages
        scratchpad_attachment_read(attachment_id="...", page_range="1,3,5")
        
        # Read pages 6-10 (after reading 1-5)
        scratchpad_attachment_read(attachment_id="...", page_range="6-10")
        
        # Read all pages (for small documents)
        scratchpad_attachment_read(attachment_id="...", page_range="all")
    """
    tool_name = "scratchpad_attachment_read"
    logger.info(f"Tool called: {tool_name} with attachment_id={attachment_id}")
    
    # Get tenant_id from context (required for download URL)
    tenant_id = ctx.deps.get("tenant_id")
    if not tenant_id:
        available_keys = list(ctx.deps.keys()) if ctx.deps else []
        logger.error(f"Tool '{tool_name}' failed: tenant_id not found in context. Available keys: {available_keys}")
        raise ValueError(
            f"tenant_id is required in workflow context for downloading attachments. "
            f"Available context keys: {available_keys}"
        )
    
    logger.info(f"Tool '{tool_name}' executing with tenant_id={'present' if tenant_id else 'missing'}, extract_tables={extract_tables}, page_range={page_range}")
    
    try:
        # Check cache first
        if attachment_id in _download_cache:
            logger.debug(f"Using cached file for attachment {attachment_id}")
            file_bytes = _download_cache[attachment_id]
        else:
            # Download file
            file_bytes = await asyncio.to_thread(_download_attachment, attachment_id, tenant_id)
            _download_cache[attachment_id] = file_bytes
        
        # Get metadata (from cache or we'll need to fetch it)
        metadata = _metadata_cache.get(attachment_id, {})
        file_type = metadata.get("fileType")
        title = metadata.get("title", "Unknown")
        
        # Read file content with page range support
        read_result = await _read_file_content(
            file_bytes, 
            file_type, 
            extract_tables=extract_tables,
            page_range=page_range,
        )
        
        # Combine results
        result = {
            "attachment_id": attachment_id,
            "title": title,
            "file_type": file_type or "unknown",
            "content": read_result.get("content", ""),
            "metadata": {
                "size": len(file_bytes),
                **{k: v for k, v in read_result.items() 
                   if k not in ["content", "table_data", "chunking_info"]},
            },
        }
        
        # Add chunking info if present
        if "chunking_info" in read_result:
            result["chunking_info"] = read_result["chunking_info"]
        
        # Add table_data if extracted
        if "table_data" in read_result:
            result["table_data"] = read_result["table_data"]
        
        logger.info(f"Tool '{tool_name}' succeeded: extracted {len(result.get('content', ''))} characters")
        return result
    except Exception as e:
        logger.error(f"Tool '{tool_name}' failed with error: {type(e).__name__}: {str(e)}")
        raise


